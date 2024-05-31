import json
from docx import Document

def make_it_word():
    data: dict = {}
    with open('./question_data/question.json') as f:
        data = json.load(f)
    full_file = ''
    last_section_known = ''
    for _, value in data.items():
        if last_section_known != value['section']:
            last_section_known = value['section']
            full_file += f'\n\n {value["section"].upper()} \n\n'
        new_line = f'#{value["question_index"]} {value["question"]}\n\t{value["answer"]}\n\n'
        full_file += new_line
    with open('./text_file.txt', 'w', encoding='utf-8') as f:
        f.write(full_file)


def make_it_word_too():
    data = {}
    with open('./question_data/question.json') as f:
        data = json.load(f)

    doc = Document()

    # Initialize variables for ToC
    toc = doc.add_paragraph()
    toc.add_run("Table of Contents").bold = True
    toc.add_run("\n\n")

    last_section_known = ''
    index = []

    page_numbers = {}  # Store page numbers for each section

    for _, value in data.items():
        if last_section_known != value['section']:
            last_section_known = value['section']
            # Add section heading
            section_heading = doc.add_heading(value['section'].upper(), level=1)
            index.append((value['section'], section_heading))

            # Store the position of the section heading for later use
            page_numbers[value['section']] = len(doc.element.body)

        # Add question and answer
        doc.add_paragraph(f'{value["question_index"]}. {value["question"]}')
        doc.add_paragraph(f'\t{value["answer"]}\n')

    # Add index to ToC with placeholders for page numbers
    for section, heading in index:
        toc_run = toc.add_run(section)
        toc_run.bold = True
        toc.add_run("\t")  # Add some spacing
        # Add placeholder for page number
        toc.add_run("Page ").italic = True
        toc.add_run("[PageNumber]").italic = True

        # Update page number field
        section_start = page_numbers[section]
        page_number_field = doc.element.body[section_start].addnext(
            '<w:fldSimple w:instr="PAGE"/>'
        )

    # Save the document
    doc.save('./word_file_with_toc.docx')
